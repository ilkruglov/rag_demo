"""Document ingestion script for building the Qdrant-backed vector store."""
from __future__ import annotations

import argparse
import hashlib
import logging
import os
import shutil
import sys
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from tqdm import tqdm

from llama_index.core import Document, Settings as LlamaSettings, StorageContext, VectorStoreIndex
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams

from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument

from app.config import get_raw_documents_dir, get_settings
from app.services.table_parser import (
    parse_table,
    enrich_text_with_codes,
    ParsedChunk,
    TableType,
)
from app.services.semantic_enrichment import enrich_chunk_with_semantic_info
from app.services.index_names import VECTOR_INDEX_ID
from app.services.chat import _get_qdrant_client
from app.services.bm25_retriever import BM25Index, clear_bm25_cache

logging.basicConfig(level=logging.WARNING, format="%(levelname)s %(message)s")
LOGGER = logging.getLogger(__name__)


def _format_time(seconds: float) -> str:
    """Format seconds to human readable string."""
    if seconds < 60:
        return f"{seconds:.0f}с"
    minutes = int(seconds // 60)
    secs = int(seconds % 60)
    if minutes < 60:
        return f"{minutes}м {secs}с"
    hours = minutes // 60
    mins = minutes % 60
    return f"{hours}ч {mins}м"

SUPPORTED_DOC_EXTENSIONS = {".docx"}
SUPPORTED_PDF_EXTENSIONS = {".pdf"}
SUPPORTED_EXTENSIONS = SUPPORTED_DOC_EXTENSIONS | SUPPORTED_PDF_EXTENSIONS

CHUNK_SIZE = 1500
CHUNK_OVERLAP = 200


@dataclass
class ContentPart:
    """A part of document content with optional metadata."""
    text: str
    metadata: dict = field(default_factory=dict)
    is_table: bool = False


def _extract_docx_with_tables(path: Path) -> list[ContentPart]:
    """Extract text from DOCX including tables with smart parsing.

    Returns list of ContentPart objects, where tables are parsed with
    semantic enrichment and metadata extraction.
    """
    doc = DocxDocument(str(path))
    source_name = path.name
    content_parts = []

    # Get all elements in document order (paragraphs and tables interleaved)
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text.strip()
                    if text:
                        # Enrich paragraph text with codes if found
                        enriched_text, text_metadata = enrich_text_with_codes(text)
                        content_parts.append(ContentPart(
                            text=enriched_text,
                            metadata=text_metadata,
                            is_table=False,
                        ))
                    break
        elif element.tag.endswith('tbl'):  # Table
            for table in doc.tables:
                if table._tbl is element:
                    # Use smart table parser
                    parsed_chunks = parse_table(table, source_name)
                    for chunk in parsed_chunks:
                        content_parts.append(ContentPart(
                            text=chunk.text,
                            metadata=chunk.metadata,
                            is_table=True,
                        ))
                    break

    return content_parts


def _table_to_text(table) -> str:
    """Convert table to text with clear structure (legacy fallback)."""
    rows_text = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace("\n", " ")
            cells.append(cell_text)
        # Skip empty rows
        if any(cells):
            rows_text.append(" | ".join(cells))

    if not rows_text:
        return ""

    return "\n".join(rows_text)


class _DocxWithTablesDoc:
    """Simple document-like object for consistency with langchain loaders."""
    def __init__(self, page_content: str, metadata: dict):
        self.page_content = page_content
        self.metadata = metadata


import re as _re

# Кэш оглавлений документов: source_name -> {номер_раздела: название}
_TOC_CACHE: dict[str, dict[str, str]] = {}


def _parse_toc(text: str) -> dict[str, str]:
    """Парсит оглавление документа и возвращает словарь {номер: название}.

    Форматы записей в оглавлении:
    - "1. ОБЩИЕ ПОЛОЖЕНИЯ…………………	6"
    - "3.1.	Депозитарные операции	29"
    - "3.4.1.	Общий порядок открытия разделов	71"
    """
    toc = {}
    lines = text.split("\n")
    in_toc = False

    for line in lines:
        stripped = line.strip()

        # Начало оглавления
        if stripped.upper() in ("СОДЕРЖАНИЕ", "ОГЛАВЛЕНИЕ"):
            in_toc = True
            continue

        if not in_toc:
            continue

        # Конец оглавления - строка без номера после нескольких записей
        if len(toc) > 5 and stripped and not _re.match(r"^\d+\.[\d.]*", stripped):
            # Проверяем что это не продолжение длинного названия
            if not any(stripped.startswith(c) for c in ["ё", "Ё", "и", "И", "в", "В"]):
                break

        # Парсим строку оглавления: "3.4.1.	Название раздела	71"
        # Паттерн: номер (цифры с точками), затем пробел/таб, затем название
        match = _re.match(r"^(\d+\.[\d.]*)\s*[.\t ]\s*(.+?)(?:[….…\t]+\s*\d+)?$", stripped)
        if match:
            section_num = match.group(1).rstrip(".")  # "3.4.1" без конечной точки
            section_title = match.group(2).strip()
            # Убираем точки в конце названия
            section_title = _re.sub(r"[….…]+$", "", section_title).strip()
            if section_title:
                toc[section_num] = f"{section_num}. {section_title}"

        # Ограничение на размер оглавления
        if len(toc) > 200:
            break

    return toc


def _get_toc_for_document(source_name: str, full_text: str) -> dict[str, str]:
    """Возвращает оглавление документа (из кэша или парсит)."""
    if source_name in _TOC_CACHE:
        return _TOC_CACHE[source_name]

    toc = _parse_toc(full_text)
    _TOC_CACHE[source_name] = toc
    return toc


def _find_section_from_toc(text: str, toc: dict[str, str]) -> str | None:
    """Ищет номер раздела в начале текста и возвращает название из оглавления."""
    if not toc:
        return None

    first_lines = "\n".join(text.splitlines()[:5])

    # Ищем номер раздела в начале текста
    for line in first_lines.splitlines():
        stripped = line.strip()
        # Паттерн номера: "3.4.1." или "3.4.1" в начале строки
        match = _re.match(r"^(\d+\.[\d.]*)", stripped)
        if match:
            section_num = match.group(1).rstrip(".")
            if section_num in toc:
                return toc[section_num]

    return None


# Паттерны структурных элементов документов НРД
_SECTION_PATTERNS = [
    # Структурные элементы документа
    _re.compile(r'^(Статья\s+\d+[\.\d]*\.?\s*.*)$', _re.IGNORECASE),
    _re.compile(r'^(Глава\s+[IVXLCDM\d]+[\.\d]*\.?\s*.*)$', _re.IGNORECASE),
    _re.compile(r'^(Раздел\s+[IVXLCDM\d]+[\.\d]*\.?\s*.*)$', _re.IGNORECASE),
    _re.compile(r'^(ЧАСТЬ\s+[IVXLCDM]+\.?\s*.*)$', _re.IGNORECASE),  # "ЧАСТЬ I ОБЩИЕ ПОЛОЖЕНИЯ"
    _re.compile(r'^(Приложение\s*[№#]?\s*\d+.*)$', _re.IGNORECASE),

    # Нумерованные разделы с табуляцией (основной формат НРД)
    _re.compile(r'^(\d+\.\t.{5,120})$'),  # "1.\tОбщие положения"
    _re.compile(r'^(\d+\.\d+\.\t.{5,120})$'),  # "1.1.\tПодраздел"
    _re.compile(r'^(\d+\.\d+\.\d+\.\t.{5,120})$'),  # "1.1.1.\tПодподраздел"

    # Нумерованные разделы с пробелом
    _re.compile(r'^(\d+\.\s+[А-ЯЁA-Z].{5,120})$'),  # "1. Общие положения"
    _re.compile(r'^(\d+\.\d+\.?\s+[А-ЯЁA-Z].{5,120})$'),  # "1.1. Подраздел" или "1.1 Подраздел"
    _re.compile(r'^(\d+\.\d+\.\d+\.?\s+[А-ЯЁA-Z].{5,100})$'),  # "1.1.1. Глубокий подраздел"

    # Римские цифры
    _re.compile(r'^([IVXLCDM]+\.\s+[А-ЯЁA-Z].{5,80})$'),  # "I. Раздел"

    # НРД-специфичные форматы
    _re.compile(r'^(Форма\s+[A-ZА-ЯЁ]{1,3}\d{2,4}.*)$', _re.IGNORECASE),  # "Форма MF035", "Форма GF034"
    _re.compile(r'^(Порядок\s+заполнения\s+.{10,100})$', _re.IGNORECASE),  # "Порядок заполнения Поручения..."
    _re.compile(r'^(Правила\s+заполнения\s+.{10,100})$', _re.IGNORECASE),  # "Правила заполнения..."
    _re.compile(r'^(Таблица\s*[№#]?\s*\d+.*)$', _re.IGNORECASE),  # "Таблица № 1"
    _re.compile(r'^(Перечень\s+.{10,100})$', _re.IGNORECASE),  # "Перечень документов..."

    # Оглавление и содержание
    _re.compile(r'^(СОДЕРЖАНИЕ|Оглавление)$', _re.IGNORECASE),

    # Термины и определения (частый раздел)
    _re.compile(r'^(Термины\s+и\s+определения)$', _re.IGNORECASE),
    _re.compile(r'^(Общие\s+положения)$', _re.IGNORECASE),

    # Коды форм как заголовки (AA001, AF005 и т.д.)
    _re.compile(r'^([A-Z]{2}\d{3})$'),  # "AA001", "AF005"
]

# Человекочитаемые названия документов по префиксам файлов
_DOCUMENT_TITLES = {
    'poryadok_nrd': 'Порядок взаимодействия НРД',
    'poryadok_p1': 'Порядок НРД - Приложение 1',
    'poryadok_p2': 'Порядок НРД - Приложение 2 (Рыночная стоимость)',
    'poryadok_p3': 'Порядок НРД - Приложение 3',
    'poryadok_p4': 'Порядок НРД - Приложение 4',
    'poryadok_p5': 'Порядок НРД - Приложение 5',
    'poryadok_p6': 'Порядок НРД - Приложение 6',
    'poryadok_p9': 'Порядок НРД - Приложение 9',
    'pravila_clearing': 'Правила клиринга',
    'pravila_dep': 'Правила депозитарной деятельности',
    'por_prov_den': 'Порядок проведения денежных расчетов',
    'por_reg_clear': 'Порядок регистрации клиринга',
    'clearing_doc': 'Документы клиринга',
    'usl_cd': 'Условия осуществления депозитарной деятельности',
}


def _get_human_document_title(source_name: str | None) -> str | None:
    """Return human-readable document title based on filename."""
    if not source_name:
        return None
    stem = Path(source_name).stem.lower()
    for prefix, title in _DOCUMENT_TITLES.items():
        if stem.startswith(prefix):
            return title
    return None


def _extract_section_title(text: str) -> str | None:
    """Return likely section heading from chunk text using structural patterns."""
    first_lines = "\n".join(text.splitlines()[:15])

    for line in first_lines.splitlines():
        candidate = line.strip()
        if not candidate or len(candidate) > 200:
            continue

        # Проверяем структурные паттерны
        for pattern in _SECTION_PATTERNS:
            match = pattern.match(candidate)
            if match:
                return match.group(1).strip()

        # Fallback: UPPERCASE заголовки (короткие, 5-80 символов)
        if candidate.isupper() and 5 < len(candidate) < 80:
            return candidate

    return None


def load_document(path: Path, pbar=None):
    suffix = path.suffix.lower()
    if suffix == ".docx":
        if pbar:
            pbar.set_postfix_str(f"DOCX: {path.name[:30]}")
        try:
            content_parts = _extract_docx_with_tables(path)
            if not content_parts:
                return []
            # Create separate documents for each content part, preserving metadata
            documents = []
            for part in content_parts:
                metadata = {"source": str(path)}
                # Copy metadata from table parser (operation_codes, form_codes, content_type)
                if part.metadata:
                    metadata.update(part.metadata)
                if part.is_table:
                    metadata["is_table"] = True
                documents.append(_DocxWithTablesDoc(
                    page_content=part.text,
                    metadata=metadata,
                ))
            return documents
        except Exception as exc:  # noqa: BLE001
            LOGGER.warning("Failed to load DOCX %s: %s", path, exc)
            return []
    elif suffix in SUPPORTED_PDF_EXTENSIONS:
        if pbar:
            pbar.set_postfix_str(f"PDF: {path.name[:30]}")
        try:
            loader = PyPDFLoader(str(path))
            return loader.load()
        except Exception as exc:  # noqa: BLE001
            LOGGER.warning("Failed to load PDF %s: %s", path, exc)
            return []
    else:
        return []


def load_documents(source: Path | Iterable[Path]):
    if isinstance(source, Path) and source.is_dir():
        paths = [path for path in source.rglob("*") if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS]
    elif isinstance(source, Path):
        paths = [source]
    else:
        paths = list(source)

    print(f"\n{'='*60}")
    print(f"  ИНДЕКСАЦИЯ ДОКУМЕНТОВ")
    print(f"{'='*60}")
    print(f"\n[1/4] Загрузка документов ({len(paths)} файлов)")
    sys.stdout.flush()

    documents = []
    with tqdm(paths, desc="Загрузка", unit="файл",
              bar_format='{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}]',
              dynamic_ncols=True, file=sys.stdout) as pbar:
        for path in pbar:
            documents.extend(load_document(path, pbar))

    print(f"  ✓ Загружено страниц: {len(documents)}")
    return documents


def _infer_category(source_name: str | None) -> str | None:
    if not source_name:
        return None
    stem = Path(source_name).stem
    if not stem:
        return None
    parts = stem.split("_")
    if not parts:
        return None
    prefix = parts[0]
    if prefix.isalpha() and len(prefix) <= 20:
        return prefix.lower()
    return stem.lower()


def normalize_documents(documents):
    normalized = []
    seen = set()
    for doc in documents:
        lines = [line.strip() for line in doc.page_content.splitlines() if line.strip()]
        page_content = "\n".join(lines)
        if not page_content:
            continue
        metadata = dict(doc.metadata)
        source = metadata.get("source")
        if source:
            source_name = Path(source).name
            metadata["source"] = source_name
            document_title = Path(source).stem
            metadata.setdefault("document_title", document_title)
            category = _infer_category(source_name)
            if category:
                metadata.setdefault("document_category", category)
        else:
            metadata.setdefault("document_title", "document")

        digest = hashlib.sha1(page_content.encode("utf-8"), usedforsecurity=False).hexdigest()
        fingerprint = (metadata.get("source"), metadata.get("page"), digest)
        if fingerprint in seen:
            continue
        seen.add(fingerprint)
        normalized.append((page_content, metadata))
    return normalized


def build_nodes(
    normalized_docs,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
    domain_rules_enabled: bool = False,
):
    splitter = SentenceSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
    print(f"\n[2/4] Разбиение на чанки ({len(normalized_docs)} страниц, chunk={chunk_size})")
    sys.stdout.flush()

    # Собираем полный текст документов для парсинга оглавления
    doc_texts: dict[str, str] = {}
    for text, metadata in normalized_docs:
        source = metadata.get("source", "")
        if source:
            if source in doc_texts:
                doc_texts[source] += "\n" + text
            else:
                doc_texts[source] = text

    # Парсим оглавления для всех документов
    print("  Парсинг оглавлений...")
    toc_count = 0
    for source, full_text in doc_texts.items():
        toc = _get_toc_for_document(source, full_text)
        if toc:
            toc_count += 1
            LOGGER.info("Parsed TOC for %s: %d sections", source, len(toc))
    print(f"  ✓ Найдено оглавлений: {toc_count} из {len(doc_texts)} документов")
    sys.stdout.flush()

    documents = [Document(text=text, metadata=metadata) for text, metadata in normalized_docs]
    nodes = splitter.get_nodes_from_documents(documents, show_progress=True)
    print(f"  ✓ Создано чанков: {len(nodes)}")

    # Статистика извлечения заголовков
    toc_matches = 0
    regex_matches = 0

    for idx, node in enumerate(nodes):
        node.metadata.setdefault("chunk_id", idx)
        source_name = node.metadata.get("source")
        if source_name:
            # Сохраняем техническое имя файла
            node.metadata["document_filename"] = Path(source_name).stem
            # Добавляем человекочитаемое название
            human_title = _get_human_document_title(source_name)
            node.metadata["document_title"] = human_title or Path(source_name).stem
        else:
            node.metadata.setdefault("document_title", "document")
            node.metadata.setdefault("document_filename", "document")

        page_label = node.metadata.get("page_label")
        if page_label is not None:
            if isinstance(page_label, str) and page_label.isdigit():
                node.metadata["page"] = int(page_label)
            elif isinstance(page_label, int):
                node.metadata["page"] = page_label
        else:
            page_value = node.metadata.get("page")
            if isinstance(page_value, str) and page_value.isdigit():
                node.metadata["page"] = int(page_value)
            elif isinstance(page_value, int):
                pass  # already int, keep it

        # Сначала пробуем найти заголовок из оглавления
        section = None
        if source_name:
            toc = _TOC_CACHE.get(source_name, {})
            section = _find_section_from_toc(node.get_content(), toc)
            if section:
                toc_matches += 1

        # Если не нашли в оглавлении, используем regex-паттерны
        if not section:
            section = _extract_section_title(node.get_content())
            if section:
                regex_matches += 1

        if section:
            node.metadata["section_title"] = section

        # Optional domain-specific semantic enrichment.
        if domain_rules_enabled and node.metadata.get("operation_codes"):
            enriched_text, enriched_meta = enrich_chunk_with_semantic_info(
                node.get_content(),
                node.metadata
            )
            # Update metadata with direction info
            node.metadata.update(enriched_meta)

    # Count semantic enrichment stats
    direction_count = sum(1 for n in nodes if n.metadata.get("operation_direction"))

    print(f"  ✓ Заголовки из оглавления: {toc_matches}")
    print(f"  ✓ Заголовки по regex: {regex_matches}")
    print(f"  ✓ Без заголовка: {len(nodes) - toc_matches - regex_matches}")
    if domain_rules_enabled:
        print(f"  ✓ С направлением операции: {direction_count}")
    return nodes


def persist_vector_store(nodes, profile_id: str | None = None):
    settings = get_settings(profile_id)
    start_time = time.time()

    if settings.force_offline_mode:
        os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")
        os.environ.setdefault("HF_HUB_OFFLINE", "1")

    print(f"\n[3/4] Генерация эмбеддингов ({len(nodes)} чанков)")
    print(f"  Модель: {settings.embedding_model}")
    sys.stdout.flush()

    embed_model = HuggingFaceEmbedding(
        model_name=settings.embedding_model,
        query_instruction="query: ",
        text_instruction="passage: ",
    )
    LlamaSettings.embed_model = embed_model
    LlamaSettings.llm = None

    vector_dimension = len(embed_model.get_text_embedding("dimension probe"))
    print(f"  Размерность: {vector_dimension}")
    sys.stdout.flush()

    qdrant_path = Path(settings.qdrant_path)
    qdrant_path.mkdir(parents=True, exist_ok=True)

    client = _get_qdrant_client(str(qdrant_path))
    if client.collection_exists(settings.qdrant_collection):
        client.delete_collection(collection_name=settings.qdrant_collection)
    client.create_collection(
        collection_name=settings.qdrant_collection,
        vectors_config=VectorParams(size=vector_dimension, distance=Distance.COSINE),
    )

    vector_store = QdrantVectorStore(
        client=client,
        collection_name=settings.qdrant_collection,
        prefer_grpc=False,
    )

    storage_dir = Path(settings.storage_dir)
    if storage_dir.exists():
        for child in storage_dir.iterdir():
            if child.is_file():
                child.unlink()
            else:
                shutil.rmtree(child)
    storage_dir.mkdir(parents=True, exist_ok=True)

    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    embed_start = time.time()
    vector_index = VectorStoreIndex(nodes, storage_context=storage_context, show_progress=True)
    vector_index.set_index_id(VECTOR_INDEX_ID)
    embed_time = time.time() - embed_start
    print(f"  ✓ Эмбеддинги созданы за {_format_time(embed_time)}")
    sys.stdout.flush()

    print(f"\n[4/4] Создание BM25 индекса с лемматизацией")
    sys.stdout.flush()
    bm25_start = time.time()
    bm25_index = BM25Index(nodes)
    bm25_index_path = storage_dir / "bm25_index.json"
    bm25_index.save(bm25_index_path)
    bm25_time = time.time() - bm25_start
    print(f"  ✓ BM25 индекс создан за {_format_time(bm25_time)}")

    print("  Сохранение индексов...")
    sys.stdout.flush()
    storage_context.persist(persist_dir=str(storage_dir))

    total_time = time.time() - start_time
    print(f"\n{'='*60}")
    print(f"  ✓ ИНДЕКСАЦИЯ ЗАВЕРШЕНА!")
    print(f"{'='*60}")
    print(f"  Всего чанков:        {len(nodes)}")
    print(f"  Размерность:         {vector_dimension}")
    print(f"  Время выполнения:    {_format_time(total_time)}")
    print(f"  Векторное хранилище: {settings.qdrant_path}")
    print(f"  Хранилище индексов:  {storage_dir}")
    print(f"{'='*60}\n")


def _clear_directory_contents(path: Path) -> None:
    if not path.exists():
        return
    for child in path.iterdir():
        if child.is_file():
            child.unlink()
        else:
            shutil.rmtree(child)


def clear_indexes(profile_id: str | None = None) -> None:
    """Remove vector and keyword indexes from local storage."""

    settings = get_settings(profile_id)
    qdrant_path = Path(settings.qdrant_path)
    qdrant_path.mkdir(parents=True, exist_ok=True)
    client = _get_qdrant_client(str(qdrant_path))

    if client.collection_exists(settings.qdrant_collection):
        client.delete_collection(collection_name=settings.qdrant_collection)

    storage_dir = Path(settings.storage_dir)
    _clear_directory_contents(storage_dir)
    storage_dir.mkdir(parents=True, exist_ok=True)
    clear_bm25_cache()


def ingest(source: Path | Iterable[Path] | None = None, profile_id: str | None = None):
    settings = get_settings(profile_id)
    raw_dir = get_raw_documents_dir(settings)

    if source is None:
        source = raw_dir

    documents = load_documents(source)
    if not documents:
        raise RuntimeError("No documents found for ingestion")

    normalized = normalize_documents(documents)
    chunk_size = getattr(settings, "ingest_chunk_size", CHUNK_SIZE)
    chunk_overlap = getattr(settings, "ingest_chunk_overlap", CHUNK_OVERLAP)

    if chunk_size < 100:
        raise ValueError(f"chunk_size must be at least 100, got {chunk_size}")
    if chunk_overlap < 0:
        raise ValueError(f"chunk_overlap must be non-negative, got {chunk_overlap}")
    if chunk_overlap >= chunk_size:
        raise ValueError(f"chunk_overlap ({chunk_overlap}) must be less than chunk_size ({chunk_size})")

    nodes = build_nodes(
        normalized,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        domain_rules_enabled=settings.domain_rules_enabled,
    )
    persist_vector_store(nodes, profile_id=profile_id)


def main(raw_dir: Path | None = None, profile_id: str | None = None):
    if raw_dir is not None and not raw_dir.exists():
        msg = f"Source directory not found: {raw_dir}"
        raise FileNotFoundError(msg)

    ingest(raw_dir, profile_id=profile_id)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Ingest documents into the vector store")
    parser.add_argument("--source", type=Path, default=None, help="Directory with raw documents")
    parser.add_argument("--profile", type=str, default=None, help="Task profile identifier")
    args = parser.parse_args()
    main(args.source, profile_id=args.profile)
